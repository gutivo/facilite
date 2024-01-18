# views.py
from django.shortcuts import render
from django.http import HttpResponse
from django.views import View
from django.views.generic import TemplateView
from .forms import SeuFormulario
from docx import Document
import os
from django.conf import settings
from django.shortcuts import get_object_or_404
from cadastros.models import Campo
import pandas as pd 
import datetime
import locale
from django.contrib.auth.mixins import LoginRequiredMixin
from braces.views import GroupRequiredMixin
from django.urls import reverse_lazy
from usuarios.models import Perfil
import requests 
from django.http import JsonResponse
from docx.shared import Inches
# Configurando a localização para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')

class SuaVisualizacao(GroupRequiredMixin,LoginRequiredMixin,TemplateView):
    login_url = reverse_lazy('login')
    group_required = [u"administrador",u"colaborador", u"cliente"]
    template_name = 'formabrir.html'
    modelo_path = os.path.join(settings.BASE_DIR, 'static', 'docs', 'disp', 'MODELO05.docx')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)


        # Obtém o perfil associado ao usuário logado
        perfil = get_object_or_404(Perfil, usuario=self.request.user)

        # Adiciona a URL da foto ao contexto
        context['foto'] = perfil.foto.url if perfil.foto else ''
        
        return context  
    
    def get(self, request):
        

        form = SeuFormulario()
        
         # Chama get_context_data para obter o contexto
        context = self.get_context_data()

        # Adiciona o formulário ao contexto
        context['form'] = form

        # Renderiza a página com o contexto completo
        return render(request, self.template_name, context)
    

    def post(self, request):  
        form = SeuFormulario(request.POST)
        if form.is_valid():
            dados = form.cleaned_data
            tipo_de_poder_escolhido = form.cleaned_data['poder']
            tipo_de_gen_escolhido = form.cleaned_data['gen']
            tipo_de_pros_escolhido = form.cleaned_data['descr2']
            tipo_de_pros_DISINEX = form.cleaned_data['tipopro']
            tipo_de_pros_INCISO = form.cleaned_data['tipoinciso']
            campo_instance = get_object_or_404(Campo, id=dados['campo'].id)
            # Supondo que você tenha a ID da instância do modelo Campo
            campo_instance2 =get_object_or_404(Campo, foto=campo_instance.foto)

            # Obtendo o caminho do arquivo da imagem
            image_path = campo_instance2.foto.path
            cnpj = form.cleaned_data['cnpjx']  # Obtém o CNPJ da solicitação GET
            rs = None
            cnpjx = None
            SF1 = None
            if not cnpj:
             return JsonResponse({'error': 'CNPJ não fornecido'}, status=400)

            url = f'https://www.receitaws.com.br/v1/cnpj/{cnpj}'

            
            response = requests.get(url)

            if response.status_code == 200:
                   data = response.json()

                # Extraindo os dados necessários
                   nome = data.get('nome', '')
                   cnpjj = data.get('cnpj', '')
                   rs = nome
                   cnpjx = cnpjj
                   representante_legal = data.get('nome_rep_legal', '')
                   atividade_principal = data.get('atividade_principal', [{'text': 'Nenhum registro encontrado', 'code': ''}])[0]
                   atividades_secundarias = data.get('atividades_secundarias', [{'text': 'Nenhum registro encontrado', 'code': ''}])[0]
           

            # Criar o documento Word
            document = Document(self.modelo_path)
            # Adicione um cabeçalho ao documento
            header = document.sections[0].header
            header.paragraphs[0].add_run().add_picture(image_path, width=Inches(2.0))

            # Substituir as tags no documento
            for paragrafo in document.paragraphs:
                for run in paragrafo.runs:
                    if tipo_de_poder_escolhido == 'PODER LEGISLATIVO':
                        run.text = run.text.replace('[TIPOPODER2]', 'da Câmara')
                        run.text = run.text.replace('[GESTOR]', campo_instance.prefpresi)
                        run.text = run.text.replace('[CARGOGESTOR]', 'PRESIDENTE')
                        if tipo_de_gen_escolhido == 'MASCULINO':
                            run.text = run.text.replace('[TRATAMENTO]', 'O Presidente da Câmara Municipal')
                            run.text = run.text.replace('[TRATAMENTOS]', 'Excelentíssimo Senhor')
                            run.text = run.text.replace('[TRATAMENTO02]', 'Prezado Senhor')
                            run.text = run.text.replace('[TIPOGESTOR]', 'do Presidente do')
                            run.text = run.text.replace('[TIPOGESTOR2]', 'O Presidente')
                            run.text = run.text.replace('[TIPOGESTOR3]', 'do Presidente')
                        elif tipo_de_gen_escolhido == 'FEMININO':

                            run.text = run.text.replace('[TRATAMENTO]', 'A Presidente da Câmara Municipal')
                            run.text = run.text.replace('[TRATAMENTOS]', 'Excelentíssima Senhora')
                            run.text = run.text.replace('[TRATAMENTO02]', 'Prezada Senhora')
                            run.text = run.text.replace('[TIPOGESTOR]', 'da Presidente da')
                            run.text = run.text.replace('[TIPOGESTOR2]', 'A Presidente')
                            run.text = run.text.replace('[TIPOGESTOR3]', 'da Presidente')
                    else:
                        run.text = run.text.replace('[TIPOPODER2]', 'do Fundo')
                        if tipo_de_gen_escolhido == 'MASCULINO':
                            run.text = run.text.replace('[TRATAMENTO]', 'O Gestor Municipal')
                            run.text = run.text.replace('[TRATAMENTOS]', 'Excelentíssimo Senhor')
                            run.text = run.text.replace('[TRATAMENTO02]', 'Prezado Senhor')
                            run.text = run.text.replace('[TIPOGESTOR]', 'do gestor do')
                            run.text = run.text.replace('[TIPOGESTOR2]', 'O Gestor')
                            run.text = run.text.replace('[TIPOGESTOR3]', 'do gestor')
                        elif tipo_de_gen_escolhido == 'FEMININO':
                            run.text = run.text.replace('[TRATAMENTO]', 'A Gestora Municipal')
                            run.text = run.text.replace('[TRATAMENTOS]', 'Excelentíssima Senhora')
                            run.text = run.text.replace('[TRATAMENTO02]', 'Prezada Senhora')
                            run.text = run.text.replace('[TIPOGESTOR]', 'da gestora do')
                            run.text = run.text.replace('[TIPOGESTOR2]', 'A Gestora')
                            run.text = run.text.replace('[TIPOGESTOR3]', 'da gestor') 
                        if tipo_de_poder_escolhido == 'PODER EXECUTIVO': 
                            run.text = run.text.replace('[GESTOR]', campo_instance.gestor1)
                            run.text = run.text.replace('[CARGOGESTOR]', 'GESTOR MUNICIPAL')
                        if tipo_de_poder_escolhido == 'FUNDO MUNICIPAL DE SAÚDE': 
                            run.text = run.text.replace('[GESTOR]', campo_instance.gestor2)
                            run.text = run.text.replace('[CARGOGESTOR]', 'GESTOR DO FMS') 
                        if tipo_de_poder_escolhido == 'FUNDO MUNICIPAL DE ASSISTÊNCIA SOCIAL': 
                            run.text = run.text.replace('[GESTOR]', campo_instance.gestor3)
                            run.text = run.text.replace('[CARGOGESTOR]', 'GESTOR DO FMAS')  
                        if tipo_de_poder_escolhido == 'FUNDO MUNICIPAL DO MEIO AMBIENTE': 
                            run.text = run.text.replace('[GESTOR]', campo_instance.gestor4)
                            run.text = run.text.replace('[CARGOGESTOR]', 'GESTOR DO FMMA')  

                    run.text = run.text.replace('[VTM]', dados['vtm'] ) 
                    run.text = run.text.replace('[VT]', dados['vt'] )         
                    run.text = run.text.replace('[RSOCIAL]', rs)
                    run.text = run.text.replace('[CNPJ]', cnpjx)                                  
                    run.text = run.text.replace('[CIDADE]', campo_instance.city)       
                    run.text = run.text.replace('[OBJ01]', dados['obj1'])
                    run.text = run.text.replace('[d1]', dados['descr1'])
                    run.text = run.text.replace('[SOLICITACAO]', campo_instance.agenteSolicita)
                    run.text = run.text.replace('*PESSOA03*', campo_instance.agenteCPL)
                    run.text = run.text.replace('*PESSOA04*', campo_instance.agenteContabilidade)
                    run.text = run.text.replace('*PESSOA05*', campo_instance.agenteCI)
                    run.text = run.text.replace('[NUPROS]', dados['nupros'])
                    run.text = run.text.replace('*NUPROS22*', dados['nuproto'])
                    run.text = run.text.replace('[ANOPROS]', dados['anopro'])
                    run.text = run.text.replace('[PC]', dados['pc'])
                    run.text = run.text.replace('[DOT]', dados['dot'])
                    run.text = run.text.replace('[FONT]', dados['fonte'])
                    run.text = run.text.replace('*DIA01*', dados['data1'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA02*', dados['data2'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA03*', dados['data3'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA04*', dados['data4'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA05*', dados['data5'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA06*', dados['data6'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA07*', dados['data7'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA08*', dados['data8'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA09*', dados['data9'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('*DIA10*', dados['data10'].strftime('%d de %B de %Y'))
                    run.text = run.text.replace('[FUNDOS1]', tipo_de_poder_escolhido)
                    run.text = run.text.replace('[FUNDAECONOMICIDADE]', 'Art. 5º Na aplicação desta Lei, serão observados os princípios da legalidade, da impessoalidade, da moralidade, da publicidade, da eficiência, do interesse público, da probidade administrativa, da igualdade, do planejamento, da transparência, da eficácia, da segregação de funções, da motivação, da vinculação ao edital, do julgamento objetivo, da segurança jurídica, da razoabilidade, da competitividade, da proporcionalidade, da celeridade, da economicidade e do desenvolvimento nacional sustentável, assim como as disposições do Decreto-Lei nº 4.657, de 4 de setembro de 1942 (Lei de Introdução às Normas do Direito Brasileiro)')
                    if tipo_de_poder_escolhido == 'Prestação de serviço':  
                       run.text = run.text.replace('[TIPOPROS]', 'de prestação de serviço')
                       run.text = run.text.replace('*o1*', 'serviço')
                       run.text = run.text.replace('[ORDSA]', 'SERVIÇO')
                    else:
                          run.text = run.text.replace('[TIPOPROS]', 'de aquisição')
                          run.text = run.text.replace('*o1*', 'fornecimento')
                          run.text = run.text.replace('[ORDSA]', 'FORNECIMENTO')
                    if tipo_de_pros_DISINEX == 'Dispensa':
                       run.text = run.text.replace('[TIPO]', 'DISPENSA DE LICITAÇÃO')
                       run.text = run.text.replace('[DESCRIFUNDAMENTO]', 'Art. 75. É dispensável a licitação:')
                       run.text = run.text.replace('[FUNDAMENTO02]', '75, a regra pela qual fica dispensada a prévia licitação, ora em razão de flagrante excepcionalidade, onde a licitação, em tese, seria possível, mas pela particularidade do caso, o interesse público a reputaria inconveniente, como é o caso da dispensa.')
                       if tipo_de_pros_INCISO == 'I':
                           run.text = run.text.replace("[CC]", "I - para contratação que envolva valores inferiores a R$ 114.416,65 (cento e quatorze mil quatrocentos e dezesseis reais e sessenta e cinco centavos), no caso de obras e serviços de engenharia ou de serviços de manutenção de veículos automotores; (Vide Decreto nº 11.317, de 2022)")
                       if tipo_de_pros_INCISO == 'II':
                           run.text = run.text.replace("[CC]", "II - para contratação que envolva valores inferiores a R$ 57.208,33 (cinquenta e sete mil duzentos e oito reais e trinta e três centavos), no caso de outros serviços e compras;(Vide Decreto nº 11.317, de 2022)")  
                       if tipo_de_pros_INCISO == 'III':
                           run.text = run.text.replace("[CC]", "III - para contratação que mantenha todas as condições definidas em edital de licitação realizada há menos de 1 (um) ano, quando se verificar que naquela licitação:")
                       if tipo_de_pros_INCISO == 'III, linear A':
                           run.text = run.text.replace("[CC]", "III - para contratação que mantenha todas as condições definidas em edital de licitação realizada há menos de 1 (um) ano, quando se verificar que naquela licitação:\na) não surgiram licitantes interessados ou não foram apresentadas propostas válidas;") 
                       if tipo_de_pros_INCISO == 'III, linear B':
                           run.text = run.text.replace("[CC]", "III - para contratação que mantenha todas as condições definidas em edital de licitação realizada há menos de 1 (um) ano, quando se verificar que naquela licitação:\nb) as propostas apresentadas consignaram preços manifestamente superiores aos praticados no mercado ou incompatíveis com os fixados pelos órgãos oficiais competentes;")             
                    else:
                        run.text = run.text.replace('[TIPO]', 'INEXIGIBILIDADE')
                        run.text = run.text.replace('[DESCRIFUNDAMENTO]', 'Art. 74. É inexigível a licitação quando inviável a competição, em especial nos casos de:')
                        run.text = run.text.replace('[FUNDAMENTO02]', '74, que permite a execução do processo administrativo por meio de inexigibilidade de licitação. Nesse caso, não há concorrência viável devido à singularidade do objeto a ser contratado. A contratação direta se mostra a opção mais vantajosa em termos de prazo, custo e qualidade. Assim, a inexigibilidade é justificada para garantir a eficiência e o atendimento das necessidades da administração.')
                        if tipo_de_pros_INCISO == 'I':
                           run.text = run.text.replace("[CC]", "I - aquisição de materiais, de equipamentos ou de gêneros ou contratação de serviços que só possam ser fornecidos por produtor, empresa ou representante comercial exclusivos;")
                        if tipo_de_pros_INCISO == 'II':
                           run.text = run.text.replace("[CC]", "II - contratação de profissional do setor artístico, diretamente ou por meio de empresário exclusivo, desde que consagrado pela crítica especializada ou pela opinião pública;")
                        if tipo_de_pros_INCISO == 'III':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:")
                        if tipo_de_pros_INCISO == 'III, linear A':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\na) estudos técnicos, planejamentos, projetos básicos ou projetos executivos;")
                        if tipo_de_pros_INCISO == 'III, linear B':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\nb) pareceres, perícias e avaliações em geral;")   
                        if tipo_de_pros_INCISO == 'III, linear C':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\nc) assessorias ou consultorias técnicas e auditorias financeiras ou tributárias;")  
                        if tipo_de_pros_INCISO == 'III, linear D':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\nd) fiscalização, supervisão ou gerenciamento de obras ou serviços;")  
                        if tipo_de_pros_INCISO == 'III, linear E':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\ne) patrocínio ou defesa de causas judiciais ou administrativas;") 
                        if tipo_de_pros_INCISO == 'III, linear F':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\nf) treinamento e aperfeiçoamento de pessoal;")     
                        if tipo_de_pros_INCISO == 'III, linear G':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\ng) restauração de obras de arte e de bens de valor histórico;")
                        if tipo_de_pros_INCISO == 'III, linear H':
                           run.text = run.text.replace("[CC]", "III - contratação dos seguintes serviços técnicos especializados de natureza predominantemente intelectual com profissionais ou empresas de notória especialização, vedada a inexigibilidade para serviços de publicidade e divulgação:\nh) controles de qualidade e tecnológico, análises, testes e ensaios de campo e laboratoriais, instrumentação e monitoramento de parâmetros específicos de obras e do meio ambiente e demais serviços de engenharia que se enquadrem no disposto neste inciso;")
                        if tipo_de_pros_INCISO == 'IV':
                           run.text = run.text.replace("[CC]", "IV - objetos que devam ou possam ser contratados por meio de credenciamento;")
                        if tipo_de_pros_INCISO == 'V':
                           run.text = run.text.replace("[CC]", "V - aquisição ou locação de imóvel cujas características de instalações e de localização tornem necessária sua escolha.")         
        tabela = document.tables[0]  # Assumindo que a tabela está na primeira posição
        df = None
        if 'excel_file' in request.FILES:
    # Carregar dados do arquivo Excel
         excel_file = request.FILES['excel_file']
         df = pd.read_excel(excel_file, header=None)  # Sem cabeçalho/nomes de coluna
    # Adicionar as linhas à tabela
         if excel_file != None:
          for indice, linha in df.iterrows():
           nova_linha = tabela.add_row().cells
        # Verificar se há pelo menos 3 células na linha antes de acessar os índices
           if len(nova_linha):
            # Preencher a nova linha com dados do arquivo Excel
             nova_linha[0].text = str(indice + 1)  # Adiciona 1 para começar a contar do índice 1
             nova_linha[1].text = str(linha.iloc[0])  # Acessar a primeira célula da linha
             nova_linha[2].text = str(linha.iloc[1])  # Acessar a segunda célula da linha
             nova_linha[3].text = str(linha.iloc[2])  # Acessar a terceira célula da linha
        else:
         for indice in range(1, len(request.POST)//4 ):
            nova_linha = tabela.add_row().cells
    # Verificar se há pelo menos 6 elementos na tupla antes de acessar os índices
            if len(nova_linha):
        # Preencher a nova linha com dados do formulário
              nova_linha[0].text = str(indice)
              nova_linha[1].text = request.POST.get(f'unidade_{indice}', '')
              nova_linha[2].text = request.POST.get(f'quantidade_{indice}', '')
              nova_linha[3].text = request.POST.get(f'descricao_{indice}', '')
         
#SEGUNDA TABELA,
        tabela2 = document.tables[1]  # Assumindo que a tabela está na primeira posição
        df2 = None
        if 'excel_file' in request.FILES:
    # Carregar dados do arquivo Excel
         excel_file2 = request.FILES['excel_file']
         df2 = pd.read_excel(excel_file2, header=None)  # Sem cabeçalho/nomes de coluna
    # Adicionar as linhas à tabela
         if excel_file2 != None:
          for indice, linha in df2.iterrows():
           nova_linha2 = tabela2.add_row().cells
        # Verificar se há pelo menos 3 células na linha antes de acessar os índices
           if len(nova_linha2):
            # Preencher a nova linha com dados do arquivo Excel
             nova_linha2[0].text = str(indice + 1)  # Adiciona 1 para começar a contar do índice 1
             nova_linha2[1].text = str(linha.iloc[0])  # Acessar a primeira célula da linha
             nova_linha2[2].text = str(linha.iloc[1])  # Acessar a segunda célula da linha
             nova_linha2[3].text = str(linha.iloc[2])  # Acessar a terceira célula da linha
        else:
         for indice in range(1, len(request.POST)//4 ):
            nova_linha2 = tabela2.add_row().cells
    # Verificar se há pelo menos 6 elementos na tupla antes de acessar os índices
            if len(nova_linha2):
        # Preencher a nova linha com dados do formulário
              nova_linha2[0].text = str(indice)
              nova_linha2[1].text = request.POST.get(f'unidade_{indice}', '')
              nova_linha2[2].text = request.POST.get(f'quantidade_{indice}', '')
              nova_linha2[3].text = request.POST.get(f'descricao_{indice}', '')


              #terceira TABELA,
        tabela3 = document.tables[2]  # Assumindo que a tabela está na primeira posição
        df3 = None
        if 'excel_file3' in request.FILES:
    # Carregar dados do arquivo Excel
         excel_file3 = request.FILES['excel_file3']
         df3 = pd.read_excel(excel_file3, header=None)  # Sem cabeçalho/nomes de coluna
    # Adicionar as linhas à tabela
         if excel_file3 != None:
          for indice, linha in df3.iterrows():
           nova_linha3 = tabela3.add_row().cells
        # Verificar se há pelo menos 3 células na linha antes de acessar os índices
           if len(nova_linha3):
            # Preencher a nova linha com dados do arquivo Excel
             nova_linha3[0].text = str(indice + 1)  # Adiciona 1 para começar a contar do índice 1
             nova_linha3[1].text = str(linha.iloc[0])  # Acessar a primeira célula da linha
             nova_linha3[2].text = str(linha.iloc[1])  # Acessar a segunda célula da linha
             
        else:
         for indice in range(1, len(request.POST)//3 ):
            nova_linha3 = tabela3.add_row().cells
    # Verificar se há pelo menos 6 elementos na tupla antes de acessar os índices
            if len(nova_linha3):
        # Preencher a nova linha com dados do formulário
              nova_linha3[0].text = str(indice)
              nova_linha3[1].text = request.POST.get(f'empresas_{indice}', '')
              nova_linha3[2].text = request.POST.get(f'empvt_{indice}', '')


        #quarta TABELA,
        tabela4 = document.tables[3]  # Assumindo que a tabela está na primeira posição
        df4 = None
        if 'excel_file4' in request.FILES:
    # Carregar dados do arquivo Excel
         excel_file4 = request.FILES['excel_file4']
         df4 = pd.read_excel(excel_file4, header=None)  # Sem cabeçalho/nomes de coluna
    # Adicionar as linhas à tabela
         if excel_file4 != None:
          # Variáveis para armazenar a soma
          soma_ultima_coluna = 0
          for indice, linha in df4.iterrows():
           nova_linha4 = tabela4.add_row().cells
        # Verificar se há pelo menos 3 células na linha antes de acessar os índices
           if len(nova_linha4):
            # Preencher a nova linha com dados do arquivo Excel
             nova_linha4[0].text = str(indice + 1)  # Adiciona 1 para começar a contar do índice 1
             nova_linha4[1].text = str(linha.iloc[0])  # Acessar a primeira célula da linha
             nova_linha4[2].text = str(linha.iloc[1])  # Acessar a segunda célula da linha
             nova_linha4[3].text = str(linha.iloc[2])  # Acessar a terceira célula da linha
             nova_linha4[4].text = str(linha.iloc[3])  # Acessar a terceira célula da linha
             nova_linha4[5].text = str(linha.iloc[4])  # Acessar a terceira célula da linha
             # Somar o valor do último índice de cada linha
             soma_ultima_coluna += linha.iloc[-1]
             SF1 = soma_ultima_coluna
        else:
         # Variável para armazenar a soma
         soma_ultima_coluna_formulario = 0
         for indice in range(1, len(request.POST)//6 +1 ):
            nova_linha4 = tabela4.add_row().cells
    # Verificar se há pelo menos 6 elementos na tupla antes de acessar os índices
            if len(nova_linha4) >=6:
        # Preencher a nova linha com dados do formulário
              nova_linha4[0].text = str(indice)
              nova_linha4[1].text = request.POST.get(f'und_{indice}', '')
              nova_linha4[2].text = request.POST.get(f'qnt_{indice}', '')
              nova_linha4[3].text = request.POST.get(f'descr_{indice}', '')
              nova_linha4[4].text = request.POST.get(f'param_{indice}', '')
              nova_linha4[5].text = request.POST.get(f'vt_{indice}', '')
            # Acessar o último valor de cada linha do formulário
              ultimo_valor = request.POST.get(f'vt_{indice}', '')

            # Verificar se o último valor é um número antes de somar
            if ultimo_valor and ultimo_valor.isdigit():
               soma_ultima_coluna_formulario += int(ultimo_valor) 

       
            
        caminho_copia = os.path.join(settings.BASE_DIR, 'copia_gerada.docx')
        document.save(caminho_copia)

        with open(caminho_copia, 'rb') as docx_file:
            response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=copia_gerada.docx'
            return response

        return render(request, self.template_name, {'form': form})
